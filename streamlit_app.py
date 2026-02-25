import streamlit as st
import requests
import json
from pathlib import Path
from io import BytesIO
from PIL import Image

# Mirror TEXT_EXTENSIONS from backend
TEXT_EXTENSIONS = {
    '.txt', '.rtf', '.log', '.md', '.csv', '.json', '.xml',
    '.yaml', '.yml', '.py', '.js', '.html', '.css', '.c',
    '.cpp', '.java', '.sh', '.odt', '.tex', '.rst'
}

# Load API configuration
try:    
    API_BASE = st.secrets["api_base"]
    API_KEY = st.secrets.get("api_key", "")
except Exception:
    API_BASE = "http://localhost:8000"
    API_KEY = ""


def inject_css():
    css = """
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Poppins:wght@300;400;600;700&display=swap');
    
    /* Background: Off-white, Text: Black, Accents: Oil Pink to Wine Red */
    [data-testid='stAppViewContainer'] { 
            font-family: 'Poppins', sans-serif; 
            font-size:16px;
            background: #f5f3f0 !important;
            color: #1a1a1a !important;
    }
    
    /* Remove shadows from Streamlit navbar/header */
    [data-testid='stHeader'] {
            box-shadow: none !important;
            filter: drop-shadow(none) !important;
    }
    
    header {
            box-shadow: none !important;
            filter: drop-shadow(none) !important;
    }
    
    /* Remove shadows from top navigation */
    nav {
            box-shadow: none !important;
            filter: drop-shadow(none) !important;
    }
    
    /* Text color overrides */
    body, p, span, div, h1, h2, h3, h4, h5, h6 {
            color: #1a1a1a !important;
    }
    
    .app-header { display:flex; align-items:center; gap:20px; margin-bottom:12px; }
    .app-title { font-size:32px; font-weight:700; margin:0; color: #8b3a50; }
    .app-sub { color:#5a4a55; margin-top:6px; font-size:14px; }

    /* Oil Pink to Wine Red palette */
    :root { --accent: #a85572; --accent-soft: rgba(168,85,114,0.12); --accent-strong: #c85a73; }

    /* Enhanced skill badges with hover */
    .skill-badge { 
            display:inline-block; 
            background:var(--accent-soft); 
            color:#8b3a50; 
            padding:6px 10px; 
            border-radius:999px; 
            margin:6px 8px 6px 0; 
            font-weight:600; 
            font-size:14px;
            transition: all 0.2s ease;
            cursor: default;
            border: 1px solid rgba(168,85,114,0.3);
    }
    .skill-badge:hover { 
            background: rgba(168,85,114,0.25); 
            transform: translateY(-1px);
            box-shadow: 0 4px 8px rgba(168,85,114,0.3);
    }
    
    /* Enhanced card with better shadow */
    .card { 
            background: #fdfcfb;
            border:1px solid rgba(168,85,114,0.2); 
            border-radius:12px; 
            padding:18px; 
            box-shadow: 0 2px 8px rgba(168,85,114,0.1);
            margin-bottom: 16px;
            color: #1a1a1a;
    }
    
    .muted { color: #6a5a65 }
    
    /* Section spacing */
    .section-header { margin-top: 20px !important; margin-bottom: 12px !important; }

    /* Buttons */
    .primary-btn > button { background-color:var(--accent) !important; color:white !important; border-radius:8px !important; padding:10px 16px !important; transition: all 0.2s ease !important; }
    .primary-btn > button:hover { background-color:#8b3a50 !important; }
    .download-btn > button { background-color: #8b3a50 !important; color:white !important; border-radius:8px !important; padding:10px 16px !important; transition: all 0.2s ease !important; }
    .download-btn > button:hover { background-color:#6e2d40 !important; }

    /* Headings & text sizes */
    h1, .app-title { font-size:32px; color: #8b3a50; }
    h2 { font-size:20px; color: #1a1a1a; }
    .card small, .muted { font-size:13px }
    
    /* Empty state styling */
    .empty-state {
            text-align: center;
            padding: 48px 24px;
            background: rgba(168,85,114,0.06);
            border-radius: 12px;
            border: 1px solid rgba(168,85,114,0.15);
            color: #b3597d;
    }
    .empty-state-emoji { font-size: 48px; margin-bottom: 16px; }
    
    /* Error panel */
    .error-panel {
            background: rgba(200,90,115,0.1);
            border: 1px solid rgba(200,90,115,0.3);
            border-radius: 8px;
            padding: 16px;
            margin: 12px 0;
            color: #8b3a50;
    }

    /* Align columns to top and normalize spacing */
    .stColumns { align-items: flex-start !important; }
    .stColumns > div { padding-top: 0px !important; margin-top: 0px !important; }
    h1, h2, h3, .app-title { margin-top: 0px !important; margin-bottom: 8px !important; }
    [data-testid="stFileUploader"], [data-testid="fileUploader"] { margin-top: 0px !important; padding-top: 0px !important; }
    
    /* Hover effect on expanders */
    [data-testid="stExpander"] { 
            transition: all 0.2s ease;
    }
    [data-testid="stExpander"]:hover { 
            border-color: rgba(168,85,114,0.3) !important;
    }

    /* Tabs: subtle, modern minimalist design */

    [role="tablist"] { 
            border-bottom: 1px solid rgba(168,85,114,0.15) !important;
            box-shadow: none !important;
            padding-bottom: 0 !important;
            gap: 2px !important;
            filter: drop-shadow(none) !important;
    }
    
    /* Remove shadow from tab container */
    [data-testid="stTabs"] {
            box-shadow: none !important;
            filter: drop-shadow(none) !important;
    }

    [role="tab"] {
            padding: 10px 14px !important;
            margin-right: 0 !important;
            border-radius: 6px !important;
            border: none !important;
            background: transparent !important;
            transition: all 0.2s ease;
            font-size: 15px !important;
            font-weight: 500 !important;
            letter-spacing: 0.4px !important;
            color: #5a4a55 !important;
            box-shadow: none !important;
            filter: drop-shadow(none) !important;
    }

    [role="tab"]:hover { 
            background: rgba(168,85,114,0.1) !important; 
            cursor: pointer; 
            color: #8b3a50 !important;
            box-shadow: none !important;
            filter: drop-shadow(none) !important;
    }

    [role="tab"][aria-selected="true"] {
            background: rgba(168,85,114,0.15) !important;
            color: #8b3a50 !important;
            border-bottom: 2px solid #8b3a50 !important;
            box-shadow: none !important;
            filter: drop-shadow(none) !important;
    }
    
    /* Remove shadow from tab panels */
    [role="tabpanel"] {
            box-shadow: none !important;
            filter: drop-shadow(none) !important;
    }

    /* Feather-style inline icons */
    .feather-icon svg { stroke: #8b3a50; width: 18px; height: 18px; vertical-align: middle; margin-right: 8px; }
    .tab-heading { display:flex; align-items:center; gap:8px; margin-bottom:8px; }
    .feather-small svg { width:16px; height:16px; margin-right:6px; vertical-align:middle; stroke: #8b3a50; }
    
    /* File Uploader Styling */
    [data-testid="stFileUploader"] {
            background: #fdfcfb !important;
            border: 2px solid #c85a73 !important;
            border-radius: 12px !important;
            padding: 20px !important;
    }
    
    [data-testid="stFileUploader"] label {
            color: #1a1a1a !important;
            font-weight: 600 !important;
    }
    
    /* Upload area styling */
    [data-testid="stFileUploaderDropzone"] {
            background: rgba(168,85,114,0.08) !important;
            border: none !important;
            border-radius: 8px !important;
    }
    
    .uploadedFile {
            background: rgba(168,85,114,0.05) !important;
            border: 1px solid #c85a73 !important;
            color: #1a1a1a !important;
    }
    
    /* Remove button styling for uploaded files */
    .uploadedFile button {
            background-color: transparent !important;
            color: #8b3a50 !important;
            border: none !important;
            box-shadow: none !important;
    }
    
    .uploadedFile button:hover {
            background-color: transparent !important;
            color: #6e2d40 !important;
    }

    /* Make file-uploader pagination arrows transparent ("<" and ">") */
    [data-testid="stFileUploader"] button {
            background-color: transparent !important;
            color: #8b3a50 !important;
            border: none !important;
            box-shadow: none !important;
    }
    [data-testid="stFileUploader"] button:hover {
            background-color: transparent !important;
            color: #6e2d40 !important;
    }
    
    /* Browse files button */
    [data-testid="stFileUploaderDropzone"] button {
            background-color: #c85a73 !important;
            color: white !important;
            border-radius: 6px !important;
            font-weight: 600 !important;
    }
    
    [data-testid="stFileUploaderDropzone"] button:hover {
            background-color: #8b3a50 !important;
    }
    
    /* Parse button styling */
    button {
            background-color: #eddade !important;
            color: white !important;
            border-radius: 8px !important;
            font-weight: 800 !important;
            border: 2px solid #c85a73 !important;
            transition: all 0.3s ease !important;
    }
    
    button:hover {
            background-color: #9e3c50 !important;
            box-shadow: #1a1a1a !important;
    }
    
    button:active {
            background-color: #ffffff !important;
    }
    
    /* Selectbox/Dropdown styling */
    [data-testid="stSelectbox"] label {
            color: #b3597d !important;
    }
    
    [data-testid="stSelectbox"] div[data-baseweb="select"] > div {
            background-color: #fdfcfb !important;
            border: 1px solid #c85a73 !important;
            color: #b3597d !important;
            border-radius: 8px !important;
    }

    [data-testid="stSelectbox"] div[data-baseweb="select"] > div:hover {
            border-color: #8b3a50 !important;
            background-color: rgba(168,85,114,0.05) !important;
    }
    
    /* Dropdown menu items */
    div[data-baseweb="popover"] {
            background-color: #ffffff !important;
    }
    
    div[data-baseweb="menu"] {
            background-color: #ffffff !important;
    }
    
    div[data-baseweb="menu"] li {
            background-color: #ffffff !important;
            color: #1a1a1a !important;
    }
    
    div[data-baseweb="menu"] li:hover, div[data-baseweb="menu"] li[aria-selected="true"] {
            background-color: rgba(168,85,114,0.15) !important;
            color: #1a1a1a !important;
    }
    
    /* Additional dropdown options styling */
    [role="option"] {
            background-color: #ffffff !important;
            color: #1a1a1a !important;
    }
    
    [role="option"]:hover {
            background-color: rgba(168,85,114,0.15) !important;
            color: #1a1a1a !important;
    }
    
    /* Listbox styling */
    [role="listbox"] {
            background-color: #ffffff !important;
    }
    
    /* Force white background on all dropdown-related elements */
    [data-baseweb="popover"] * {
            background-color: inherit !important;
    }
    
    [data-baseweb="select"] [role="listbox"] {
            background-color: #ffffff !important;
    }
    
    [data-baseweb="select"] [role="option"] {
            background-color: #ffffff !important;
            color: #1a1a1a !important;
    }
    
    /* SVG icon color in dropdown */
    [data-testid="stSelectbox"] svg {
            fill: #8b3a50 !important;
    }
    
    [data-testid="stSelectbox"] > div > div > button:hover {
            background-color: rgba(168,85,114,0.08) !important;
            border-color: rgba(168,85,114,0.5) !important;
    }
    
    /* Dropdown menu styling */
    [data-testid="stSelectbox"] ul {
            background-color: #ffffff !important;
    }
    
    [data-testid="stSelectbox"] li {
            color: #1a1a1a !important;
            background-color: #ffffff !important;
    }
    
    [data-testid="stSelectbox"] li:hover {
            background-color: rgba(168,85,114,0.15) !important;
            color: #1a1a1a !important;
    }
    

    /* Progress bar styling - Pink theme */
    [data-testid="stProgress"] > div > div > div > div {
            background-color: #c85a73 !important;
    }
    
    [data-testid="stProgress"] > div > div > div {
            background-color: rgba(168,85,114,0.2) !important;
    }
    
    /* JSON viewer styling */
    [data-testid="stJsonLiteral"] {
            background-color: #2d2d2d !important;
            border-radius: 8px !important;
            padding: 16px !important;
    }

    [data-testid="stJsonLiteral"] * {
            background-color: transparent !important;
            color: #e0e0e0 !important;
    }

    /* JSON keys and values */
    .json-key {
            color: #88c0d0 !important;
    }

    .json-value {
            color: #a3be8c !important;
    }

    /* JSON brackets and punctuation */
    [data-testid="stJsonLiteral"] span {
            color: #eceff4 !important;
    }

    /* Value color styling */
    .data-value {
            color: #a85572 !important;
            font-weight: 500;
    }

    /* Text Area styling - black text for resume content */
    [data-testid="stTextArea"] textarea {
            color: #030303 !important;
            background-color: #fdfcfb !important;
            font-family: 'Courier New', monospace !important;
    }

    [data-testid="stTextArea"] label {
            color: #030303 !important;
            font-weight: 600;
    }

    </style>
    """
    st.markdown(css, unsafe_allow_html=True)


def post_file_to_endpoint(endpoint: str, file_bytes: bytes, filename: str):
    url = f"{API_BASE.rstrip('/')}/{endpoint.lstrip('/')}"
    files = {'file': (filename, file_bytes)}
    headers = {}
    # Always use API key from secrets/config
    try:
        secret_key = st.secrets.get('api_key', '')
    except Exception:
        secret_key = ''
    if secret_key:
        headers['X-API-Key'] = secret_key
    try:
        resp = requests.post(url, files=files, headers=headers, timeout=180)
        return resp
    except Exception as e:
        st.error(f"Request failed: {e}")
        return None
ICON_USER = """
    <svg xmlns='http://www.w3.org/2000/svg' viewBox='0 0 24 24' fill='none' stroke='currentColor' stroke-width='2' stroke-linecap='round' stroke-linejoin='round' class='feather feather-user'><path d='M20 21v-2a4 4 0 0 0-4-4H8a4 4 0 0 0-4 4v2'></path><circle cx='12' cy='7' r='4'></circle></svg>
    """
ICON_ID = """
    <svg xmlns='http://www.w3.org/2000/svg' viewBox='0 0 24 24' fill='none' stroke='currentColor' stroke-width='2' stroke-linecap='round' stroke-linejoin='round' class='feather feather-credit-card'><rect x='2' y='5' width='20' height='14' rx='2' ry='2'></rect><line x1='2' y1='10' x2='22' y2='10'></line></svg>
    """
ICON_MAIL = """
    <svg xmlns='http://www.w3.org/2000/svg' viewBox='0 0 24 24' fill='none' stroke='currentColor' stroke-width='2' stroke-linecap='round' stroke-linejoin='round' class='feather feather-mail'><path d='M4 4h16v16H4z'></path><polyline points='22,6 12,13 2,6'></polyline></svg>
    """
ICON_PHONE = """
    <svg xmlns='http://www.w3.org/2000/svg' viewBox='0 0 24 24' fill='none' stroke='currentColor' stroke-width='2' stroke-linecap='round' stroke-linejoin='round' class='feather feather-phone'><path d='M22 16.92v3a2 2 0 0 1-2.18 2 19.86 19.86 0 0 1-8.63-3.07 19.5 19.5 0 0 1-6-6 19.86 19.86 0 0 1-3.07-8.67A2 2 0 0 1 4.11 2h3a2 2 0 0 1 2 1.72c.12 1.21.36 2.4.72 3.53a2 2 0 0 1-.45 2.11L8.09 11.91a16 16 0 0 0 6 6l1.54-1.54a2 2 0 0 1 2.11-.45c1.13.36 2.32.6 3.53.72A2 2 0 0 1 22 16.92z'></path></svg>
    """
ICON_PIN = """
    <svg xmlns='http://www.w3.org/2000/svg' viewBox='0 0 24 24' fill='none' stroke='currentColor' stroke-width='2' stroke-linecap='round' stroke-linejoin='round' class='feather feather-map-pin'><path d='M21 10c0 7-9 11-9 11S3 17 3 10a9 9 0 0 1 18 0z'></path><circle cx='12' cy='10' r='3'></circle></svg>
    """
ICON_FILE = """
    <svg xmlns='http://www.w3.org/2000/svg' viewBox='0 0 24 24' fill='none' stroke='currentColor' stroke-width='2' stroke-linecap='round' stroke-linejoin='round' class='feather feather-file-text'><path d='M14 2H6a2 2 0 0 0-2 2v16a2 2 0 0 0 2 2h12a2 2 0 0 0 2-2V8z'></path><polyline points='14 2 14 8 20 8'></polyline><line x1='16' y1='13' x2='8' y2='13'></line><line x1='16' y1='17' x2='8' y2='17'></line><line x1='10' y1='9' x2='8' y2='9'></line></svg>
    """
ICON_STAR = """
    <svg xmlns='http://www.w3.org/2000/svg' viewBox='0 0 24 24' fill='none' stroke='currentColor' stroke-width='2' stroke-linecap='round' stroke-linejoin='round' class='feather feather-star'><polygon points='12 2 15 9 22 9 17 14 19 21 12 17 5 21 7 14 2 9 9 9 12 2'></polygon></svg>
    """
ICON_GLOBE = """
    <svg xmlns='http://www.w3.org/2000/svg' viewBox='0 0 24 24' fill='none' stroke='currentColor' stroke-width='2' stroke-linecap='round' stroke-linejoin='round' class='feather feather-globe'><circle cx='12' cy='12' r='10'></circle><line x1='2' y1='12' x2='22' y2='12'></line><path d='M12 2a15.3 15.3 0 0 1 4 10 15.3 15.3 0 0 1-4 10 15.3 15.3 0 0 1-4-10A15.3 15.3 0 0 1 12 2z'></path></svg>
    """
ICON_BRIEFCASE = """
    <svg xmlns='http://www.w3.org/2000/svg' viewBox='0 0 24 24' fill='none' stroke='currentColor' stroke-width='2' stroke-linecap='round' stroke-linejoin='round' class='feather feather-briefcase'><rect x='2' y='7' width='20' height='14' rx='2' ry='2'></rect><path d='M16 3h-8v4h8V3z'></path></svg>
    """
ICON_BOOK = """
    <svg xmlns='http://www.w3.org/2000/svg' viewBox='0 0 24 24' fill='none' stroke='currentColor' stroke-width='2' stroke-linecap='round' stroke-linejoin='round' class='feather feather-book'><path d='M4 19.5A2.5 2.5 0 0 1 6.5 17H20'></path><path d='M4 4h16v13H4z'></path></svg>
    """
ICON_DOWNLOAD = """
    <svg xmlns='http://www.w3.org/2000/svg' viewBox='0 0 24 24' fill='none' stroke='currentColor' stroke-width='2' stroke-linecap='round' stroke-linejoin='round' class='feather feather-download'><path d='M21 15v4a2 2 0 0 1-2 2H5a2 2 0 0 1-2-2v-4'></path><polyline points='7 10 12 15 17 10'></polyline><line x1='12' y1='15' x2='12' y2='3'></line></svg>
    """
ICON_X_CIRCLE = """
    <svg xmlns='http://www.w3.org/2000/svg' viewBox='0 0 24 24' fill='none' stroke='currentColor' stroke-width='2' stroke-linecap='round' stroke-linejoin='round' class='feather feather-x-circle'><circle cx='12' cy='12' r='10'></circle><line x1='15' y1='9' x2='9' y2='15'></line><line x1='9' y1='9' x2='15' y2='15'></line></svg>
    """
ICON_SEARCH = """
    <svg xmlns='http://www.w3.org/2000/svg' viewBox='0 0 24 24' fill='none' stroke='currentColor' stroke-width='2' stroke-linecap='round' stroke-linejoin='round' class='feather feather-search'><circle cx='11' cy='11' r='7'></circle><line x1='21' y1='21' x2='16.65' y2='16.65'></line></svg>
    """
ICON_CHECK_CIRCLE = """
    <svg xmlns='http://www.w3.org/2000/svg' viewBox='0 0 24 24' fill='none' stroke='currentColor' stroke-width='2' stroke-linecap='round' stroke-linejoin='round' class='feather feather-check-circle'><path d='M22 11.08V12a10 10 0 1 1-5.93-9.14'></path><polyline points='22 4 12 14.01 9 11.01'></polyline></svg>
    """


def render_skills(skills):
    if not skills:
            return
    html = ""
    for s in skills:
            html += f"<span class='skill-badge'>{s}</span>"
    st.markdown(html, unsafe_allow_html=True)


def render_contact_row(data):
    # Align all BIO fields in a single column with clear indexing
    bio_fields = []
    name = data.get('name', '')
    if name:
            bio_fields.append(("Name", name))
    email = data.get('email')
    if email:
            bio_fields.append(("Email", email))
    phone = data.get('phone')
    if phone:
            bio_fields.append(("Phone", phone))
    addr = data.get('address') or {}
    addr_str = ", ".join([v for v in [addr.get('street_address'), addr.get('district'), addr.get('state'), addr.get('country')] if v])
    if addr_str:
            bio_fields.append(("Location", addr_str))

    for label, value in bio_fields:
            # Use a flex row for icon and text
            if label == 'Name':
                st.markdown(f"<div class='contact-row'><span class='feather-small'>{ICON_ID}</span> <strong>{label}:</strong> <span class='data-value'>{value}</span></div>", unsafe_allow_html=True)
            elif label == 'Email':
                st.markdown(f"<div class='contact-row'><span class='feather-small'>{ICON_MAIL}</span> <strong>{label}:</strong> <span class='data-value'>{value}</span></div>", unsafe_allow_html=True)
            elif label == 'Phone':
                st.markdown(f"<div class='contact-row'><span class='feather-small'>{ICON_PHONE}</span> <strong>{label}:</strong> <span class='data-value'>{value}</span></div>", unsafe_allow_html=True)
            elif label == 'Location':
                st.markdown(f"<div class='contact-row'><span class='feather-small'>{ICON_PIN}</span> <strong>{label}:</strong> <span class='data-value'>{value}</span></div>", unsafe_allow_html=True)
            else:
                st.markdown(f"<div class='contact-row'><strong>{label}:</strong> <span class='data-value'>{value}</span></div>")


def pretty_render(result, file_id=None, file_bytes=None, original_filename=None):
    data = result if result else {}
    with st.container():
            tabs = st.tabs([
                "Bio", 
                "Description", 
                "Skills", 
                "Languages", 
                "Employment History", 
                "Qualifications",
                "Original Resume"
            ])
            # BIO tab
            with tabs[0]:
                st.markdown(f"<div class='tab-heading'><span class='feather-icon'>{ICON_USER}</span><strong>BIO</strong></div>", unsafe_allow_html=True)
                render_contact_row(data)
                st.markdown("</div>", unsafe_allow_html=True)
            # Description tab
            with tabs[1]:
                st.markdown(f"<div class='tab-heading'><span class='feather-icon'>{ICON_FILE}</span><strong>Description</strong></div>", unsafe_allow_html=True)
                summary = data.get('summary')
                if summary:
                    st.markdown(f"<b>Summary:</b><br>{summary}", unsafe_allow_html=True)
                else:
                    st.markdown("<span class='muted'>No summary provided</span>", unsafe_allow_html=True)
                st.markdown("</div>", unsafe_allow_html=True)
            # Skills tab
            with tabs[2]:
                st.markdown(f"<div class='tab-heading'><span class='feather-icon'>{ICON_STAR}</span><strong>Skills</strong></div>", unsafe_allow_html=True)
                render_skills(data.get('skills', []))
                st.markdown("</div>", unsafe_allow_html=True)
            # Languages tab
            with tabs[3]:
                st.markdown(f"<div class='tab-heading'><span class='feather-icon'>{ICON_GLOBE}</span><strong>Languages</strong></div>", unsafe_allow_html=True)
                langs = data.get('languages', [])
                if langs:
                    for l in langs:
                        st.markdown(f"- {l.get('language')} (read: {l.get('can_read')}, speak: {l.get('can_speak')}, write: {l.get('can_write')})")
                else:
                    st.markdown("<span class='muted'>No languages specified</span>", unsafe_allow_html=True)
                st.markdown("</div>", unsafe_allow_html=True)
            # Employment tab
            with tabs[4]:
                st.markdown(f"<div class='tab-heading'><span class='feather-icon'>{ICON_BRIEFCASE}</span><strong>Employment History</strong></div>", unsafe_allow_html=True)
                emp_list = data.get('employment', [])
                if emp_list:
                    for idx, emp in enumerate(emp_list[:10], start=1):
                        with st.expander(f"{emp.get('designation') or emp.get('company_name') or 'Role'} — {emp.get('company_name') or ''}"):
                            for k, v in emp.items():
                                if v:
                                    st.markdown(f"**{k}**: {v}")
                else:
                    st.markdown("<span class='muted'>No employment history</span>", unsafe_allow_html=True)
                st.markdown("</div>", unsafe_allow_html=True)
            # Qualifications tab
            with tabs[5]:
                st.markdown(f"<div class='tab-heading'><span class='feather-icon'>{ICON_BOOK}</span><strong>Qualifications</strong></div>", unsafe_allow_html=True)
                quals = data.get('qualifications', [])
                if quals:
                    for q in quals:
                        st.markdown(f"- {q.get('qualification')} — {q.get('university_name') or ''} ({q.get('year_of_completion') or ''})")
                else:
                    st.markdown("<span class='muted'>No qualifications listed</span>", unsafe_allow_html=True)
                st.markdown("</div>", unsafe_allow_html=True)
            # Original Resume tab
            with tabs[6]:
                st.markdown(f"<div class='tab-heading'><span class='feather-icon'>{ICON_FILE}</span><strong>Original Resume</strong></div>", unsafe_allow_html=True)
                if file_bytes and original_filename:
                    file_ext = Path(original_filename).suffix.lower()
                    
                    # Display based on file type
                    if file_ext == '.pdf':
                        import base64
                        pdf_base64 = base64.b64encode(file_bytes).decode()
                        pdf_display = f'<iframe src="data:application/pdf;base64,{pdf_base64}" width="100%" height="800" type="application/pdf"></iframe>'
                        st.markdown(pdf_display, unsafe_allow_html=True)
                    elif file_ext in ['.jpg', '.jpeg', '.png', '.bmp', '.gif', '.webp']:
                        from PIL import Image
                        from io import BytesIO
                        image = Image.open(BytesIO(file_bytes))
                        st.image(image, use_column_width=True)
                    elif file_ext in ['.txt', '.md', '.log', '.csv', '.json', '.xml', '.yaml', '.yml']:
                        # Text files - decode and display
                        for encoding in ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']:
                            try:
                                text_content = file_bytes.decode(encoding)
                                st.text_area("Resume Content", text_content, height=400, disabled=True)
                                break
                            except (UnicodeDecodeError, UnicodeError):
                                continue
                    elif file_ext in ['.docx']:
                        # Word documents - extract and display content
                        try:
                            from docx import Document
                            from io import BytesIO
                            doc = Document(BytesIO(file_bytes))
                            
                            # Extract all paragraphs and tables
                            content = []
                            for para in doc.paragraphs:
                                if para.text.strip():
                                    content.append(para.text)
                            
                            # Extract tables
                            for table in doc.tables:
                                for row in table.rows:
                                    row_text = " | ".join([cell.text.strip() for cell in row.cells])
                                    if row_text.strip():
                                        content.append(row_text)
                            
                            # Display extracted content
                            full_content = "\n".join(content)
                            st.text_area("Word Document Content", full_content, height=400, disabled=True)
                        except Exception as e:
                            st.error(f"Error reading Word document: {str(e)}\n\nMake sure python-docx is installed: `pip install python-docx`")
                    elif file_ext in ['.doc']:
                        st.info(f"📄 **{original_filename}** (.doc format)\n\nOlder .doc format is not directly supported. Please use .docx format for best results. The parsed data is available in other tabs.")
                    else:
                        st.info(f"📄 **{original_filename}** (File type: {file_ext})\n\nUnable to display this file type in preview. Check other tabs for parsed data.")
                else:
                    st.markdown("<span class='muted'>No resume file available for viewing.</span>", unsafe_allow_html=True)
                st.markdown("</div>", unsafe_allow_html=True)


def main():
    # Page config must be called before any other Streamlit commands
    st.set_page_config(page_title="AI-Resume Parser — Frontend", layout='wide')
    inject_css()

    # Header
    with st.container():
            st.markdown("<div class='app-header'><div><h1 class='app-title'>Resume Parser</h1><div class='app-sub'>Upload resumes and get structured JSON</div></div></div>", unsafe_allow_html=True)

    left, right = st.columns([1.3, 2])

    with left:
            st.markdown("### Upload")
            
            # Allow single or multiple file selection
            uploaded_files = st.file_uploader(
                "Choose resume file(s) or ZIP",
                type=None,
                accept_multiple_files=True,
                help="Select files from folder (Ctrl+Click multiple) or upload a ZIP archive"
            )
            
            if uploaded_files:
                st.markdown("\n")
                
                # Handle multiple files or ZIP
                is_zip = len(uploaded_files) == 1 and Path(uploaded_files[0].name).suffix.lower() == '.zip'
                is_multiple = len(uploaded_files) > 1
                
                # Determine processing type
                button_label = "Parse All" if (is_zip or is_multiple) else "Parse"
                status_upload = "Uploading files..." if is_multiple else ("Uploading ZIP..." if is_zip else "Uploading file...")
                status_process = "Processing resumes..." if (is_zip or is_multiple) else "Processing document..."
                
                if st.button(button_label, key='parse', use_container_width=True):
                    # Process files directly with pink progress bar
                    progress_bar = st.progress(0)
                    status_text = st.empty()
                    
                    if is_zip or is_multiple:
                        # Batch processing
                        status_text.text(status_upload)
                        progress_bar.progress(30)
                        
                        if is_zip:
                            # Single ZIP file
                            file_bytes = uploaded_files[0].read()
                            progress_bar.progress(50)
                            status_text.text(status_process)
                            resp = post_file_to_endpoint('parse_batch', file_bytes, uploaded_files[0].name)
                        else:
                            # Multiple files - create ZIP
                            import zipfile
                            from io import BytesIO
                            zip_buffer = BytesIO()
                            with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
                                for uploaded_file in uploaded_files:
                                    file_bytes = uploaded_file.read()
                                    zip_file.writestr(uploaded_file.name, file_bytes)
                            zip_buffer.seek(0)
                            progress_bar.progress(50)
                            status_text.text(status_process)
                            resp = post_file_to_endpoint('parse_batch', zip_buffer.getvalue(), 'batch_upload.zip')
                        
                        progress_bar.progress(90)
                        
                        if resp and resp.status_code == 200:
                            batch_data = resp.json()
                            st.session_state['batch_results'] = batch_data
                            st.session_state['batch_mode'] = True
                            st.session_state['selected_result_index'] = 0
                            progress_bar.progress(100)
                            status_text.empty()
                            progress_bar.empty()
                            st.success(f"✅ Processed {batch_data.get('total_files', 0)} files in {batch_data.get('total_processing_time', 0):.2f}s")
                            st.rerun()
                        elif resp:
                            progress_bar.empty()
                            status_text.empty()
                            st.error(f"❌ Batch processing failed: {resp.status_code} - {resp.text}")
                        else:
                            progress_bar.empty()
                            status_text.empty()
                            st.error("❌ Request failed")
                    else:
                        # Single file processing
                        status_text.text(status_upload)
                        progress_bar.progress(30)
                        
                        file_bytes = uploaded_files[0].read()
                        filename = uploaded_files[0].name
                        
                        progress_bar.progress(50)
                        status_text.text(status_process)
                        resp = post_file_to_endpoint('parse', file_bytes, filename)
                        
                        progress_bar.progress(90)
                        
                        if resp and resp.status_code == 200:
                            result = resp.json()
                            parsed_data = result.get('data')
                            file_id = result.get('file_id')
                            processing_time = result.get('processing_time', 0)
                            
                            st.session_state['latest_parsed'] = parsed_data
                            st.session_state['latest_file_id'] = file_id
                            st.session_state['latest_file_bytes'] = file_bytes
                            st.session_state['latest_filename'] = filename
                            st.session_state['batch_mode'] = False
                            
                            progress_bar.progress(100)
                            status_text.empty()
                            progress_bar.empty()
                            st.success(f"✅ Parsed in {processing_time:.2f}s")
                            st.rerun()
                        elif resp:
                            progress_bar.empty()
                            status_text.empty()
                            st.error(f"❌ Parsing failed: {resp.status_code} - {resp.text}")
                        else:
                            progress_bar.empty()
                            status_text.empty()
                            st.error("❌ Request failed")
                        st.markdown('<h4>Authentication Required</h4>', unsafe_allow_html=True)
                        api_key_input = st.text_input('Enter API Password', type='password', key='auth_api_key')
                        submit_auth = st.form_submit_button('Authenticate')
                        if submit_auth:
                            st.session_state['api_key_input'] = api_key_input
                            st.session_state['show_auth_popup'] = False
                            # Now proceed with parse/upload logic
                            # Store file bytes in session state for batch mode
                            batch_file_bytes = {}
                            # Determine endpoint and prepare files
                            if is_multiple:
                                endpoint = 'parse_batch'
                                from io import BytesIO
                                import zipfile
                                zip_buffer = BytesIO()
                                with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zf:
                                    for f in uploaded_files:
                                        file_bytes = f.read()
                                        batch_file_bytes[f.name] = file_bytes
                                        zf.writestr(f.name, file_bytes)
                                zip_buffer.seek(0)
                                file_to_send = zip_buffer.getvalue()
                                filename_to_send = 'batch.zip'
                            elif is_zip:
                                endpoint = 'parse_batch'
                                zip_bytes = uploaded_files[0].read()
                                file_to_send = zip_bytes
                                filename_to_send = uploaded_files[0].name
                                import zipfile
                                from io import BytesIO
                                with zipfile.ZipFile(BytesIO(zip_bytes), 'r') as zf:
                                    for zip_filename in zf.namelist():
                                        if not zip_filename.endswith('/'):
                                            batch_file_bytes[zip_filename] = zf.read(zip_filename)
                            else:
                                endpoint = 'parse'
                                file_to_send = uploaded_files[0].read()
                                filename_to_send = uploaded_files[0].name
                            progress_bar = st.progress(0)
                            status_text = st.empty()
                            import time
                            start_time = time.time()
                            try:
                                status_text.markdown(f"<span style='color:#ffb3d9'>{status_upload}</span>", unsafe_allow_html=True)
                                progress_bar.progress(10 if (is_zip or is_multiple) else 33)
                                resp = post_file_to_endpoint(endpoint, file_to_send, filename_to_send)
                                status_text.markdown(f"<span style='color:#ffb3d9'>{status_process}</span>", unsafe_allow_html=True)
                                progress_bar.progress(50 if (is_zip or is_multiple) else 66)
                                if resp is None:
                                    progress_bar.empty()
                                    status_text.empty()
                                    st.markdown(f"<div class='error-panel'>{ICON_X_CIRCLE} <strong>Connection Error</strong><br/>No response from server. Please check if the backend is running.</div>", unsafe_allow_html=True)
                                    return
                                if resp.status_code == 200:
                                    status_text.markdown("<span style='color:#ffb3d9'>Extracting data...</span>", unsafe_allow_html=True)
                                    progress_bar.progress(100)
                                    try:
                                        response_data = resp.json()
                                    except Exception as e:
                                        progress_bar.empty()
                                        status_text.empty()
                                        st.markdown(f"<div class='error-panel'>{ICON_X_CIRCLE} <strong>Parse Error</strong><br/>Response is not valid JSON: {str(e)}</div>", unsafe_allow_html=True)
                                        st.code(resp.text)
                                        return
                                    time.sleep(0.3)
                                    progress_bar.empty()
                                    status_text.empty()
                                    if is_zip or is_multiple:
                                        st.session_state['batch_results'] = response_data
                                        st.session_state['batch_file_bytes'] = batch_file_bytes
                                        st.session_state['batch_mode'] = True
                                    else:
                                        parsed = response_data.get('data') or response_data
                                        file_id = response_data.get('file_id') if isinstance(response_data, dict) else None
                                        st.session_state['latest_parsed'] = parsed
                                        st.session_state['latest_file_id'] = file_id
                                        st.session_state['latest_file_bytes'] = file_to_send
                                        st.session_state['latest_filename'] = filename_to_send
                                        st.session_state['batch_mode'] = False
                                else:
                                    progress_bar.empty()
                                    status_text.empty()
                                    try:
                                        error_data = resp.json()
                                        if isinstance(error_data, dict) and 'detail' in error_data:
                                            error_detail = error_data['detail']
                                            if isinstance(error_detail, dict):
                                                error_msg = error_detail.get('message', str(error_detail))
                                            else:
                                                error_msg = str(error_detail)
                                        else:
                                            error_msg = str(error_data)
                                    except:
                                        error_msg = resp.text
                                    st.markdown(f"<div class='error-panel'>{ICON_X_CIRCLE} <strong>Server Error ({resp.status_code})</strong><br/>{error_msg}</div>", unsafe_allow_html=True)
                            except Exception as e:
                                progress_bar.empty()
                                status_text.empty()
                                st.markdown(f"<div class='error-panel'>{ICON_X_CIRCLE} <strong>Error</strong><br/>{str(e)}</div>", unsafe_allow_html=True)
                            if (is_zip or is_multiple) and st.session_state.get('batch_mode') and st.session_state.get('batch_results'):
                                st.session_state['selected_result_index'] = 0   

    with right:
            st.markdown(f"<div style='display:flex; align-items:center; gap:12px; margin-bottom:12px;'><span style='display:flex; width:36px; height:36px; min-width:36px; color:#c85a73;'>{ICON_CHECK_CIRCLE}</span><h3 style='margin:0; font-size:30px;'>Parsed Results</h3></div>", unsafe_allow_html=True)
            
            # Check if we have batch results
            if st.session_state.get('batch_mode') and st.session_state.get('batch_results'):
                batch_data = st.session_state['batch_results']
                results = batch_data.get('results', [])
                
                if results:
                    # Create dropdown for resume selection
                    st.markdown("#### Select a resume:")
                    
                    # Create options with status
                    options = []
                    for idx, result in enumerate(results):
                        filename = result['filename']
                        status = result['status']
                        status_icon = "✓" if status == "success" else "✗"
                        options.append(f"{status_icon} {filename}")
                    
                    # Dropdown selection
                    selected_option = st.selectbox(
                        "Choose a resume:",
                        options=options,
                        index=st.session_state.get('selected_result_index', 0),
                        label_visibility="collapsed"
                    )
                    
                    # Update selected index based on dropdown
                    if selected_option:
                        selected_idx = options.index(selected_option)
                        st.session_state['selected_result_index'] = selected_idx
                    
                    # Display selected resume
                    selected_idx = st.session_state.get('selected_result_index', 0)
                    if 0 <= selected_idx < len(results):
                        selected = results[selected_idx]
                        st.markdown("---")
                        st.markdown(f"#### {selected['filename']}")

                        if selected.get('status') == 'success' and selected.get('data'):
                            file_id = selected.get('file_id')
                            # For batch results, get file bytes from stored batch_file_bytes
                            batch_file_bytes = st.session_state.get('batch_file_bytes', {})
                            file_bytes = batch_file_bytes.get(selected['filename'])
                            filename = selected.get('filename')
                            pretty_render(selected['data'], file_id=file_id, file_bytes=file_bytes, original_filename=filename)
                            with st.expander('View Raw JSON'):
                                st.json(selected['data'])
                        else:
                            err = selected.get('error', 'Unknown error')
                            st.markdown(f"<div class='error-panel'>{ICON_X_CIRCLE} <strong>Error</strong><br/>{err}</div>", unsafe_allow_html=True)
            
            # Single file mode
            elif st.session_state.get('latest_parsed'):
                parsed = st.session_state['latest_parsed']
                file_id = st.session_state.get('latest_file_id')
                file_bytes = st.session_state.get('latest_file_bytes')
                filename = st.session_state.get('latest_filename')
                if parsed:
                    pretty_render(parsed, file_id=file_id, file_bytes=file_bytes, original_filename=filename)
                    with st.expander('View Raw JSON'):
                        st.json(parsed)
            
            # Default empty state
            else:
                st.markdown(f"""
                <div class='empty-state'>
                    <div class='empty-state-emoji' style='font-size:48px; margin-bottom:16px;'>
                        <span style='vertical-align:middle; margin-left:8px;'>✨</span>
                    </div>
                    <h3 style='color: #ffd6eb; margin-bottom: 8px;'>No Results Yet</h3>
                    <p class='muted'>Upload a resume file on the left and click <strong>Parse</strong> to get started</p>
                </div>
                """, unsafe_allow_html=True)


if __name__ == '__main__':
    main()
